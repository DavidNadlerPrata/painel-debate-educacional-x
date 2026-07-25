# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BRAND = RGBColor(0x2B, 0x3A, 0x8F)
INK2  = RGBColor(0x48, 0x4C, 0x60)
MUTED = RGBColor(0x7B, 0x80, 0x98)

doc = Document()

# base style
st = doc.styles['Normal']
st.font.name = 'Times New Roman'
st.font.size = Pt(11)
st.paragraph_format.space_after = Pt(6)
st.paragraph_format.line_spacing = 1.3

def set_cell_bg(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)

def eyebrow(text):
    p = doc.add_paragraph()
    r = p.add_run(text.upper()); r.font.size = Pt(8.5); r.font.color.rgb = BRAND; r.bold = True
    r.font.name = 'Consolas'
    p.paragraph_format.space_after = Pt(2)
    # letter spacing
    rPr = r._element.get_or_add_rPr(); spc = OxmlElement('w:spacing'); spc.set(qn('w:val'), '30'); rPr.append(spc)
    return p

def h1(text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(14); r.font.color.rgb = RGBColor(0x14,0x16,0x1F)
    return p

def h2(text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(11.5); r.font.color.rgb = BRAND
    return p

def body(text, italic=False, color=None, after=6):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text); r.italic = italic
    if color is not None: r.font.color.rgb = color
    return p

def bullet(text_runs):
    p = doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after = Pt(3)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if isinstance(text_runs, str): text_runs = [(text_runs, False)]
    for t, b in text_runs:
        r = p.add_run(t); r.bold = b
    return p

def rule():
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr(); pbdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'4'); bot.set(qn('w:space'),'1'); bot.set(qn('w:color'),'CFD2E0')
    pbdr.append(bot); pPr.append(pbdr)

# ---------------- COVER ----------------
eyebrow('Proposta de prova de conceito · Painel de monitoramento')
p = doc.add_paragraph(); r = p.add_run('Inteligência Social no Debate Educacional no X')
r.bold = True; r.font.size = Pt(20); r.font.color.rgb = RGBColor(0x14,0x16,0x1F)
p.paragraph_format.space_after = Pt(2)
p = doc.add_paragraph(); r = p.add_run('Da Ficha de Especificação de Alvo (FEA) ao painel: um demonstrador do método de especificação situada do alvo para detecção de posicionamento')
r.font.size = Pt(12); r.font.color.rgb = INK2
p.paragraph_format.space_after = Pt(10)

t = doc.add_table(rows=4, cols=2); t.alignment = WD_TABLE_ALIGNMENT.LEFT
meta = [('Programa','PPGGTD/UFT — Doutorado Profissional em Governança e Transformação Digital'),
        ('Doutorando','Leonardo Barchini'),
        ('Orientador','Prof. Dr. David Nadler Prata'),
        ('Vínculo na tese','Capítulo 7 — Aplicação institucional; artefato central: Ficha de Especificação de Alvo')]
for i,(k,v) in enumerate(meta):
    c0 = t.cell(i,0); c1 = t.cell(i,1)
    rr = c0.paragraphs[0].add_run(k); rr.bold=True; rr.font.size=Pt(9.5); rr.font.color.rgb=MUTED
    rr2 = c1.paragraphs[0].add_run(v); rr2.font.size=Pt(10)
    c0.width = Cm(3.6); c1.width = Cm(12)
for row in t.rows:
    for cell in row.cells:
        cell.paragraphs[0].paragraph_format.space_after = Pt(2)
rule()

# ---------------- 1 ----------------
h1('1. O que é esta PoC')
body('Esta proposta descreve uma prova de conceito (PoC) — um painel de monitoramento funcional — que materializa, de forma visual e navegável, a tese sustentada pela pesquisa: em monitoramento governamental orientado a eventos, a validade dos indicadores de posicionamento depende mais da especificação do alvo do que da escolha do classificador. O painel não é o objeto científico da tese; é o demonstrador do artefato que é — a Ficha de Especificação de Alvo (FEA) — e o veículo pelo qual a contribuição chega à equipe de comunicação do Ministério da Educação (MEC).')
body('A PoC foi construída como protótipo interativo de página única (HTML autocontido), operável em navegador, com dados sintéticos realistas. Ela pode ser exibida em banca, em oficina de transferência no MEC e como figura viva do Capítulo 7.')

p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(4)
r = p.add_run('Protótipo publicado: '); r.bold=True; r.font.size=Pt(10)
r2 = p.add_run('https://claude.ai/code/artifact/2b882966-907d-4c00-ac8e-b9677713d7d3'); r2.font.size=Pt(10); r2.font.color.rgb=BRAND

# ---------------- 2 ----------------
h1('2. A afirmação que o painel torna visível')
body('O painel foi desenhado em torno de um gesto único: mostrar a mesma multidão de manifestações lida de dois modos, lado a lado.')
bullet([('Leitura sem alvo (polaridade afetiva) — ', True), ('a prática corrente de mercado, que mede negativo/neutro/positivo sem definir sobre o que as pessoas se posicionam.', False)])
bullet([('Leitura com FEA (posicionamento vs. alvo especificado) — ', True), ('favorável / contrário / neutro / misto em relação a um alvo formulado, versionado e auditável.', False)])
body('No episódio-âncora do protótipo — o pagamento da parcela de março do Pé-de-Meia — a leitura sem alvo acusa 67% de rejeição aparente; especificado o alvo, a rejeição real ao programa cai para 20%, e o painel revela que 31% das manifestações têm alvo diferente da política nominal (deslocamento de alvo). É a reprodução, em ambiente controlado, do achado do artigo do Pé-de-Meia do grupo, em que relatos eufóricos de beneficiários foram classificados como negativos — a melhor justificativa empírica para a virada metodológica da tese.')

# ---------------- 3 ----------------
h1('3. Componentes do painel e seu vínculo com a tese')
tbl = doc.add_table(rows=1, cols=3); tbl.style = 'Table Grid'; tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = tbl.rows[0].cells
for i,txt in enumerate(['Componente do painel','O que exibe','Elemento da tese']):
    set_cell_bg(hdr[i], 'E9EBF8')
    rr = hdr[i].paragraphs[0].add_run(txt); rr.bold=True; rr.font.size=Pt(9.5); rr.font.color.rgb=BRAND
rows = [
 ('Série de menções e picos','12 meses de volume com episódios ancorados em eventos externos verificáveis','Objetivo (a); detecção de episódios por rajada (Kleinberg)'),
 ('Ficha de Especificação de Alvo','Alvo canônico, escopo dentro/fora, critérios com exemplos-âncora, alvos concorrentes descartados, analista/versão','Artefato central (DSR); objetivo (b)'),
 ('Leitura sem alvo × com FEA','Duas barras comparadas + delta do indicador de rejeição','Experimento fatorial; núcleo da tese'),
 ('Deslocamento de alvo','Decomposição da massa: programa, governo (deslocado), celebração, ruído','Objetivo (f) — quantificação inédita p/ o caso BR'),
 ('Tipologia da crítica','Operacional, distributiva, de mérito, político-partidária','Substitui a classificação ideológica (cap. 8)'),
 ('Experimento fatorial A/B/C/D','Indicador de rejeição e F1-macro por estratégia de alvo, contra o gold','Cap. 5 — avaliação do método'),
 ('Backtesting','Latência de alerta, acerto de alvo e alarme falso por crise conhecida','Cap. 5 — transforma a coleta retrospectiva em vantagem'),
]
for compn, exib, elem in rows:
    c = tbl.add_row().cells
    for j,txt in enumerate([compn, exib, elem]):
        pr = c[j].paragraphs[0]; rr = pr.add_run(txt); rr.font.size=Pt(9)
        if j==0: rr.bold=True
        pr.paragraph_format.space_after = Pt(2)
tbl.columns[0].width = Cm(3.8); tbl.columns[1].width = Cm(7.2); tbl.columns[2].width = Cm(5.2)

# ---------------- 4 ----------------
h1('4. Arquitetura da PoC (pipeline mínimo)')
for i,(t1,t2) in enumerate([
 ('Coleta','corpus retrospectivo de 12 meses do X via V-Tracker (conteúdo público, minimização e pseudonimização desde a origem).'),
 ('Detecção de episódios','baseline estatístico sobre a média móvel + detector de rajadas (Kleinberg); limiar declarado, sensibilidade e falsos alarmes documentados.'),
 ('Especificação do alvo','para cada episódio, preenchimento da FEA — por analista humano (braço C) e por LLM sob o mesmo protocolo (braço D).'),
 ('Classificação','dois classificadores mantidos constantes: um transformer em português ajustado sobre o gold standard e um LLM instruído.'),
 ('Agregação e painel','indicadores por alvo, deslocamento de alvo, tipologia da crítica; renderização no painel (a PoC atual).'),
]):
    bullet([(f'{i+1}. {t1} — ', True), (t2, False)])
body('Na PoC os cinco estágios estão representados por dados sintéticos; a versão de produção substitui cada estágio pelos dados reais do projeto MEC, sem alterar a estrutura do painel.', italic=True, color=INK2)

# ---------------- 5 ----------------
h1('5. Dados, escopo e o que a PoC deliberadamente não faz')
bullet('Fonte única: plataforma X, via V-Tracker — assegura comparabilidade metodológica entre episódios.')
bullet('Recorte: ENEM, SISU, Pé-de-Meia, ensino superior e falas do ministro, ao longo de 12 meses.')
bullet([('Fora do escopo: ', True), ('grafos de interação e detecção de comunidades, identificação de desinformação/comportamento coordenado, análise multiplataforma e qualquer perfilamento ideológico de usuários.', False)])

# ---------------- 6 ----------------
h1('6. Critérios de sucesso da PoC')
bullet([('Legibilidade institucional — ', True), ('um secretário de comunicação entende, sem treino, por que o mesmo pico produz dois números de rejeição.', False)])
bullet([('Rastreabilidade — ', True), ('todo indicador do painel reconduz a uma FEA versionada e a um evento âncora com fonte verificável.', False)])
bullet([('Fidelidade ao experimento — ', True), ('os quatro braços A/B/C/D aparecem no painel com o mesmo gold standard usado no Capítulo 5.', False)])
bullet([('Prontidão para a oficina — ', True), ('o instrumento de avaliação de utilidade (tempestividade, inteligibilidade, adequação à decisão) é aplicável sobre esta interface.', False)])

# ---------------- 7 ----------------
h1('7. Encaixe com o cronograma do projeto MEC (6 meses)')
tbl2 = doc.add_table(rows=1, cols=3); tbl2.style='Table Grid'; tbl2.alignment=WD_TABLE_ALIGNMENT.CENTER
h2c = tbl2.rows[0].cells
for i,txt in enumerate(['Mês','Produto do projeto MEC','Uso na PoC / tese']):
    set_cell_bg(h2c[i],'E9EBF8'); rr=h2c[i].paragraphs[0].add_run(txt); rr.bold=True; rr.font.size=Pt(9.5); rr.font.color.rgb=BRAND
sched = [
 ('1','Coleta retrospectiva e base estruturada','Alimenta a série e a detecção de episódios do painel (cap. 4)'),
 ('2–3','Curadoria e gold standard','Define os alvos-gold e o corpus anotado (cap. 4–5)'),
 ('3–5','Modelagem e experimento','Popula os braços A/B/C/D do painel (cap. 5)'),
 ('4–6','Análises e painel','Versão de produção da PoC (cap. 6–7)'),
 ('6','Oficinas de transferência','Avaliação de utilidade sobre a PoC (cap. 7)'),
]
for m,pr,us in sched:
    c=tbl2.add_row().cells
    for j,txt in enumerate([m,pr,us]):
        rr=c[j].paragraphs[0].add_run(txt); rr.font.size=Pt(9)
        if j==0: rr.bold=True
tbl2.columns[0].width=Cm(1.6); tbl2.columns[1].width=Cm(7.0); tbl2.columns[2].width=Cm(7.6)
body('Ressalva de prudência: o núcleo científico (cap. 5) deve ser executável sobre um subconjunto do corpus, para que a tese não fique refém do calendário do convênio.', italic=True, color=INK2)

# ---------------- 8 ----------------
h1('8. Ética e governança embutidas no painel')
body('A diferença entre escuta institucional e vigilância não é dada pela tecnologia, mas pelas regras de uso que a acompanham. A PoC as torna visíveis:')
bullet('Análise sempre agregada; vedado o perfilamento individual e qualquer inferência sobre pessoas naturais identificáveis.')
bullet('Minimização e pseudonimização desde a coleta; base legal e finalidade explícitas sob a LGPD.')
bullet('Tratamento distinto para perfis institucionais e figuras públicas; rodapé permanente sobre o que o método não é capaz de afirmar.')

# ---------------- nota ----------------
rule()
p = doc.add_paragraph(); p.paragraph_format.space_before=Pt(4)
r=p.add_run('Nota sobre o protótipo. '); r.bold=True; r.font.size=Pt(9.5)
r2=p.add_run('Todos os volumes, percentuais e fichas do painel são sintéticos, construídos para exibir o comportamento do método; não representam medições reais. A avaliação de utilidade com a equipe do MEC e o corpus anotado de referência substituem estes valores na versão final. Documento e protótipo preparados como apoio à orientação.')
r2.font.size=Pt(9.5); r2.font.color.rgb=INK2

import os
out = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Proposta_PoC_Painel_Leonardo.docx")
doc.save(out)
print("SAVED", out, os.path.getsize(out), "bytes")
